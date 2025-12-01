import streamlit as st
import tempfile
import os
import subprocess
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    PageBreak,
    Image,
    Table,
    TableStyle,
)
from reportlab.lib.pagesizes import A4, letter
    # code continues
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from PIL import Image as PILImage
import io

# ---- Streamlit page config ----
st.set_page_config(page_title="Question PDF Generator", layout="wide")
st.title("📄 Question PDF Generator")
st.markdown("Convert your questions and optional screenshots into a pro, fully-filled PDF (one question per page)")

# ---- Sidebar ----
st.sidebar.header("⚙️ Settings")
page_size_option = st.sidebar.selectbox("Page Size", ["A4", "Letter"])
page_size = A4 if page_size_option == "A4" else letter

# Slider is now for "base" size; I'll still make questions larger than this.
font_size = st.sidebar.slider("Base Question Font Size", 12, 24, 16)
show_page_numbers = st.sidebar.checkbox("Page numbers", value=True)

# ---- File uploaders ----
col1, col2 = st.columns(2)

with col1:
    st.subheader("📋 Upload Word Document")
    word_file = st.file_uploader(
        "Select .doc or .docx file with questions",
        type=["doc", "docx"],
    )

with col2:
    st.subheader("🖼️ Upload Screenshots")
    screenshot_files = st.file_uploader(
        "Upload screenshots (in order, optional)",
        type=["jpg", "jpeg", "png"],
        accept_multiple_files=True,
    )

# ---- Helper for doc -> docx conversion ----
def convert_doc_to_docx(in_path):
    try:
        output_dir = tempfile.gettempdir()
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                output_dir,
                in_path,
            ],
            check=True,
            capture_output=True,
        )
        return os.path.join(
            output_dir,
            os.path.splitext(os.path.basename(in_path))[0] + ".docx",
        )
    except Exception:
        st.error("Failed to convert .doc to .docx. Please save as .docx manually if this persists.")
        return None

# ---- Extract questions from Word file ----
questions = []
if word_file:
    is_docx = word_file.name.lower().endswith(".docx")
    with tempfile.NamedTemporaryFile(
        delete=False, suffix=(".docx" if is_docx else ".doc")
    ) as tmp_file:
        tmp_file.write(word_file.read())
        tmp_path = tmp_file.name

    if not is_docx:
        st.info("Attempting to convert .doc to .docx format...")
        tmp_path_conv = convert_doc_to_docx(tmp_path)
        if tmp_path_conv:
            tmp_path = tmp_path_conv
        else:
            st.stop()

    from docx import Document

    try:
        doc = Document(tmp_path)

        # Paragraphs:
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 5:
                questions.append(text)

        # Tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and len(text) > 5 and text not in questions:
                        questions.append(text)

        if not questions:
            st.error("No questions detected in file.")
        else:
            st.success(f"Detected {len(questions)} questions (or tasks).")
            with st.expander("Preview Questions"):
                for idx, q in enumerate(questions[:10]):
                    st.write(f"Q{idx+1}: {q[:120]}{'...' if len(q) > 120 else ''}")
                if len(questions) > 10:
                    st.write(f"...plus {len(questions) - 10} more.")
    except Exception as e:
        st.error(f"Problem reading Word document: {e}")

# ---- Screenshot mapping (purely by order: 1st -> Q1, 2nd -> Q2, ...) ----
screenshots_dict = {}  # {question_number: [UploadedFile, ...]}
if screenshot_files:
    for idx, file in enumerate(screenshot_files):
        q_index = idx + 1  # 1-based to match enumerate(questions, 1)
        if q_index not in screenshots_dict:
            screenshots_dict[q_index] = []
        screenshots_dict[q_index].append(file)

    # Optional: show mapping preview
    with st.expander("Preview Screenshot Mapping"):
        for q_num, files in screenshots_dict.items():
            st.write(f"Q{q_num} screenshots:")
            for f in files:
                st.write(f"- {f.name}")

# ---- PDF Generation ----
if questions:
    if st.button("✨ Generate Decorative PDF"):
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=page_size,
            rightMargin=0.5 * inch,
            leftMargin=0.5 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()

        # Big, bold question style
        qstyle = ParagraphStyle(
            "Question",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=font_size + 4,          # Make it larger than chosen base size
            leading=font_size + 8,           # Extra line spacing for readability
            alignment=0,                     # left
            textColor=colors.HexColor("#111111"),
        )

        # Question card wrapper style (we'll use a Table for a card effect)
        question_card_style = TableStyle(
            [
                ("BOX", (0, 0), (-1, -1), 1.2, colors.HexColor("#2c3e50")),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#eaf3ff")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )

        # Small label style (for "Screenshot:")
        label_style = ParagraphStyle(
            "Label",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=11,
            textColor=colors.HexColor("#34495e"),
            spaceAfter=0.05 * inch,
        )

        story = []

        # Usable content height (rough; for scaling screenshots)
        usable_height = page_size[1] - (0.75 * inch + 0.75 * inch)  # top + bottom margins
        # Reserve ~1/3 for question and spacing; 2/3 for screenshots
        max_screenshot_height = usable_height * 0.65

        for q_num, question in enumerate(questions, 1):
            # ----- Question card -----
            q_para = Paragraph(f"Q{q_num}: {question}", qstyle)
            q_table = Table([[q_para]], colWidths=[page_size[0] - inch])  # width ~ page width minus margins
            q_table.setStyle(question_card_style)

            story.append(q_table)
            story.append(Spacer(1, 0.25 * inch))

            # ----- Screenshots -----
            if q_num in screenshots_dict:
                story.append(Paragraph("Screenshot:", label_style))

                for uploaded_file in screenshots_dict[q_num]:
                    try:
                        uploaded_file.seek(0)
                        img_bytes = uploaded_file.read()

                        # Use PIL to get image size
                        pil_img = PILImage.open(io.BytesIO(img_bytes))
                        img_width_px, img_height_px = pil_img.size

                        # Max width ~ full content width
                        max_width = page_size[0] - inch  # full width inside margins

                        # Compute scale to fit both width and reserved height, keeping aspect ratio
                        scale_w = max_width / float(img_width_px)
                        scale_h = max_screenshot_height / float(img_height_px)
                        scale = min(scale_w, scale_h, 1.5)  # cap scale a bit so super small images don't get crazy huge

                        target_width = img_width_px * scale
                        target_height = img_height_px * scale

                        img_buffer = io.BytesIO(img_bytes)
                        img_flowable = Image(
                            img_buffer, width=target_width, height=target_height
                        )

                        # Decorative box for screenshot
                        img_table = Table([[img_flowable]], colWidths=[target_width])
                        img_table.setStyle(
                            TableStyle(
                                [
                                    ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#7f8c8d")),
                                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fbff")),
                                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                ]
                            )
                        )

                        story.append(img_table)
                        story.append(Spacer(1, 0.25 * inch))
                    except Exception as e:
                        st.warning(f"Could not embed screenshot for Q{q_num}: {e}")
            else:
                story.append(Spacer(1, 0.3 * inch))

            # Page break between questions
            if q_num < len(questions):
                story.append(PageBreak())

        # ---- Page numbers ----
        def add_page_number(canvas, doc_obj):
            if show_page_numbers:
                canvas.saveState()
                canvas.setFont("Helvetica", 9)
                canvas.setFillColor(colors.grey)
                canvas.drawCentredString(
                    page_size[0] / 2,
                    0.40 * inch,
                    f"Page {doc_obj.page}",
                )
                canvas.restoreState()

        # Build the PDF
        doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
        pdf_bytes = pdf_buffer.getvalue()

        st.success(f"Decorative PDF generated: {len(questions)} pages!")
        st.download_button(
            label="📥 Download PDF",
            data=pdf_bytes,
            file_name="questions_with_screenshots_decorative_big.pdf",
            mime="application/pdf",
        )

# ---- Info and requirements ----
with st.expander("How to use this app"):
    st.markdown(
        """
1. Upload your `.doc` or `.docx` file listing questions or tasks.
2. Optionally upload screenshots **in the same order as the questions** (1st screenshot → Q1, 2nd → Q2, etc.).
3. Click **Generate Decorative PDF**.  
   Each page = **big bold question card + large screenshot box** filling most of the page.
4. Download and print or share your PDF.

- **For `.doc` files, LibreOffice must be installed.**
- On Streamlit Cloud, add `libreoffice` to `packages.txt`.
        """
    )

st.caption("Made with Streamlit · Big, bold & decorative multipage PDF")
